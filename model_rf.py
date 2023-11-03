#Importing Important Libraries for Building RF model
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from sklearn.model_selection import cross_val_score
from sklearn.model_selection import KFold
from sklearn.model_selection import GridSearchCV
from sklearn.feature_selection import RFE
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import RobustScaler
from sklearn.ensemble import RandomForestRegressor
from sklearn import metrics
from docx import Document
from docx.shared import Inches
import warnings
warnings.filterwarnings("ignore")

# for more details about random forest 
#https://scikit-learn.org/stable/modules/generated/sklearn.ensemble.RandomForestRegressor.html
#total parameters added 7

def run(dataset,UPLOAD_FILE,parameter_1=0.3,parameter_2=0,parameter_3=100,parameter_4="mse",parameter_5=1,parameter_8="auto",parameter_12=3):
	a2 = float(parameter_1)  # Train/Test Ratio it should be float number less than 1
	a3 = bool(parameter_2)  #Normalize True/False # File CSV
	a4 = int(parameter_3) #no of estimators from 1 to positive value
	a5 = str(parameter_4) # criterion ["mse","mae"]
	a6 = int(parameter_5) # max_depth value from 1 to any positive value
	a9 = float(parameter_8) #max_features values ['auto','sqrt','log2',any positive float values in between 0 to 1 by step 0.1]
	a12 = int(parameter_12) # cross validation values range from [1,2,3,4,5]
	#dataset=pd.read_csv("nO_aCTIVITY_data_cmbl_processed.csv")
	#df=dataset
	"""
	colname=df.columns.to_list()
	print("No of Columns in dataset:\t",len(colname)+1)
	print("No of Records\t:",df.shape[0])
	print("No of Null Values\t:",df.isnull().sum().sum())
	print("No of duplicates\t:",df[df.duplicated()].shape[0])
	"""
	#Dividing dataset as features and target variable
	col=dataset.shape[1]
	nameCol=dataset.iloc[:,0:3]
	df=dataset.iloc[:,3:col]
	
	X = df.drop(['Activity'],axis=1).values
	Y = df["Activity"].values
	X_train, X_test, Y_train, Y_test = train_test_split(X, Y, test_size=a2, random_state=42)
	
	#Giving details of split
	split_details=pd.DataFrame({"X_train":[X_train.shape[0]],"X_test":[X_test.shape[0]],"Y_train":[Y_train.shape[0]],"Y_test":[Y_test.shape[0]]},index=None)
	#print("Split Details:\t"+"In ratio:",a2)
	#print(split_details)
	
	#Scaling by use of Robust Scaler
	if(a3==True):
		rc = RobustScaler()
		X_train = rc.fit_transform(X_train)
		X_test = rc.transform(X_test)
		X=rc.transform(X)
	
	#model building
	regressor_RF = RandomForestRegressor(n_estimators=a4,criterion=a5, max_depth=a6, max_features=a9)
	regressor_RF.fit(X_train, Y_train)
	y_pred_RF = regressor_RF.predict(X_test)
	y_train_RF= regressor_RF.predict(X_train)
	
	#Cross Validation
	# k-fold CV 
	RF = RandomForestRegressor(n_estimators=a4,criterion=a5, max_depth=a6, max_features=a9)
	scores = cross_val_score(RF, X_train, Y_train, scoring='r2', cv=a12)
	scores_details=pd.DataFrame({"Cross_Validation_Score":scores})
	#print("CV value:",len(scores_details))
	#print(scores_details)
	
	#Features Importance 
	feature_importance_RF=regressor_RF.feature_importances_
	feature_importance_DF=pd.DataFrame()
	feature_importance_DF["Columns"]=df.columns
	feature_importance_DF['Feature_Importance'] = pd.Series(feature_importance_RF)
	feature_importance_DF=feature_importance_DF.sort_values("Feature_Importance", ascending=False)
	#print(feature_importance_DF) 
	#Plot of first 20 important features
	plt.figure(figsize =(15, 6))
	plt.subplot(1,3,1)
	plt.title("Top20_Feature_Importance")
	sns.barplot(x='Feature_Importance' , y='Columns', data=feature_importance_DF.head(20))
	plt.savefig("savefig1_RF.png")
	
	#Errors Details
	MAE=metrics.mean_absolute_error(Y_test, y_pred_RF)
	MSE=metrics.mean_squared_error(Y_test, y_pred_RF)
	RMSE=np.sqrt(metrics.mean_squared_error(Y_test, y_pred_RF))
	"""
	print('Mean Absolute Error     :', MAE)
	print('Mean Squared Error      :', MSE)
	print('Root Mean Squared Error :', RMSE)
	"""
	R1=np.corrcoef(Y_test,y_pred_RF)[0,1]
	if (np.isnan(R1)):
		R1=0.0
	R2=np.corrcoef(Y_train,y_train_RF )[0,1]
	if (np.isnan(R2)):
		R2=0.0
	#print("Correlation Co-efficient Value on Test Set (R-Test)     :",R1)
	#print("Correlation Co-efficient Value on Test Set (R-Train)     :",R2)
	
	#Predicted Values Vs Observed_Values
	df_RF = pd.DataFrame({'Predicted_Activity_test' : y_pred_RF, 'Observed_Activity_test' : Y_test})
	#print(df_RF)
	plt.subplot(1,3,2)
	plt.title("Test_Data_Plot")
	sns.regplot(x="Observed_Activity_test", y="Predicted_Activity_test", data=df_RF)
	plt.savefig("savefig2_RF.png")
	# Full data
	y_pred_full=regressor_RF.predict(X)
	df_RF_full = pd.DataFrame({'Predicted_Activity' : y_pred_full, 'Observed_Activity' : Y})
	plt.subplot(1,3,3)
	plt.title("Full_Data_Plot")
	sns.regplot(data=df_RF_full, x="Observed_Activity", y="Predicted_Activity")
	plt.savefig(r"files/savefig3_RF.png")
	
	"""
	document = Document()
	p = document.add_paragraph()
	r = p.add_run()
	r.add_picture("savefig3_RF.png",width=Inches(7.0))
	
	p = document.add_paragraph()
	r = p.add_run()
	r.add_text('Evaluation Report of Random Forest Model')
	p = document.add_paragraph()
	r = p.add_run()
	r.add_text("Number of estimators:\t"+ "{:.2f}".format(a4))
	p = document.add_paragraph()
	r = p.add_run()
	r.add_text("Train-Test Ratio:\t"+"{:.2f}".format(a2))
	p = document.add_paragraph()
	r = p.add_run()
	r.add_text("Correlation Co-efficient Value on Test Set (R-Test): "+"{:.2f}".format(R1))
	p = document.add_paragraph()
	r = p.add_run()
	r.add_text("Correlation Co-efficient Value on Train Set (R-Train): "+"{:.2f}".format(R2))
	p = document.add_paragraph()
	r = p.add_run()
	r.add_text("Mean Absolute Error on Test Set: "+"{:.2f}".format(metrics.mean_absolute_error(Y_test, y_pred_RF)))
	p = document.add_paragraph()
	r = p.add_run()
	r.add_text("Mean Absolute Error on Train Set: "+"{:.2f}".format(metrics.mean_absolute_error(Y_train, y_train_RF)))
	p = document.add_paragraph()
	r = p.add_run()
	r.add_text("Root Mean Squared Error (RMSE) on Test Set: "+"{:.2f}".format(np.sqrt(metrics.mean_squared_error(Y_test, y_pred_RF))))
	p = document.add_paragraph()
	r = p.add_run()
	r.add_text("Root Mean Squared Error (RMSE) on Train Set: "+"{:.2f}".format(np.sqrt(metrics.mean_squared_error(Y_train, y_train_RF))))
	document.save('RF-Report.docx')
	"""
	text1 = "Correlation Co-efficient Value on Test Set (R-Test): "+"{:.2f}".format(R1)
	text2 = "Correlation Co-efficient Value on Train Set (R-Train): "+"{:.2f}".format(R2)
	text3 = "Mean Absolute Error on Test Set: "+"{:.2f}".format(metrics.mean_absolute_error(Y_test, y_pred_RF))
	text4 = "Mean Absolute Error on Train Set: "+"{:.2f}".format(metrics.mean_absolute_error(Y_train, y_train_RF))
	text5 = "Root Mean Squared Error (RMSE) on Test Set: "+"{:.2f}".format(np.sqrt(metrics.mean_squared_error(Y_test, y_pred_RF)))
	text6 = "Root Mean Squared Error (RMSE) on Train Set: "+"{:.2f}".format(np.sqrt(metrics.mean_squared_error(Y_train, y_train_RF)))
	return text1 + "\n" + text2 + "\n" + text3 + "\n" + text4 + "\n" + text5 + "\n" + text6, regressor_RF
